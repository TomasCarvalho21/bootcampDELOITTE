from src.ingestion.loaders.loaderBase import LoaderBase
from docx import Document

class LoaderDOCX(LoaderBase):

    def __init__(self, filepath:str):
        
        self.filepath=filepath

    def extract_metadata(self):

        doc = Document(self.filepath)

        # Access the core properties
        core_properties = doc.core_properties

        metadata = {  
            'author':core_properties.author,   
            'subject': core_properties.subject,  
            'title': core_properties.title,
            "Last Modified By:": core_properties.last_modified_by,
            "Created Date:": core_properties.created
        }

        self.metadata=metadata

        return self.metadata if self.all_keys_have_values(metadata=self.metadata) else False
    
    def extract_text(self):
        
        doc = Document(self.filepath)
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += "\n" + paragraph.text

        return full_text
    
    def all_keys_have_values(self, metadata, value_check=lambda x: x is not None and x != ''):
        return all(value_check(value) for value in metadata.values())
    

if __name__=="__main__":
    pass
